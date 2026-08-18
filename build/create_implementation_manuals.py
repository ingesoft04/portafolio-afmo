from pathlib import Path
from datetime import date
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "manual"
LOGO = ROOT / "public" / "fmv-logo-graphite-transparent.png"
NAVY = RGBColor(8, 43, 86)
CYAN = RGBColor(0, 174, 199)
GRAY = RGBColor(76, 86, 97)
LIGHT = "EAF7FA"


def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Calibri"
    r.font.size = Pt(9)
    if color:
        r.font.color.rgb = color
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.append(fld)


def setup(doc, lang):
    sec = doc.sections[0]
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.7)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal.font.color.rgb = GRAY
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.1
    for name, size, color in [("Title", 30, NAVY), ("Heading 1", 19, NAVY), ("Heading 2", 13.5, CYAN), ("Heading 3", 11, NAVY)]:
        s = styles[name]
        s.font.name = "Calibri"
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = color
        s.paragraph_format.space_before = Pt(10)
        s.paragraph_format.space_after = Pt(5)
        s.paragraph_format.keep_with_next = True
    header = sec.header.paragraphs[0]
    header.text = "FMV InfraSec  |  " + ("Manual de implementación" if lang == "es" else "Implementation manual")
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = GRAY
    footer = sec.footer.paragraphs[0]
    footer.add_run("FMV InfraSec  •  ")
    add_page_number(footer)


def code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.right_indent = Inches(0.15)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(8.5)
    r.font.color.rgb = NAVY
    shade_paragraph(p, "F0F4F7")


def shade_paragraph(p, fill):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def bullets(doc, items, ordered=False):
    for index, item in enumerate(items, 1):
        if ordered:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.18)
            p.add_run(f"{index}.  ").bold = True
            p.add_run(item)
        else:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(item)


def note(doc, title, body):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0, 0)
    shade(c, LIGHT)
    p = c.paragraphs[0]
    r = p.add_run(title + " — ")
    r.bold = True
    r.font.color.rgb = NAVY
    p.add_run(body)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def comparison(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        shade(t.rows[0].cells[i], "082B56")
        set_cell_text(t.rows[0].cells[i], h, True, RGBColor(255,255,255))
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val)
            if len(t.rows) % 2 == 1:
                shade(cells[i], "F4F7F9")
    doc.add_paragraph()


def cover(doc, lang):
    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO), width=Inches(2.8))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.style = doc.styles["Title"]
    p.add_run("Manual de implementación\ndel portafolio" if lang == "es" else "Portfolio implementation\nmanual")
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run("FMV InfraSec · Andrés Felipe Martínez Obando")
    r.bold = True; r.font.size = Pt(14); r.font.color.rgb = CYAN
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("Guía de publicación, actualización, dominio y seguridad" if lang == "es" else "Publishing, maintenance, domain and security guide")
    doc.add_paragraph()
    comparison(doc,
        ["Documento" if lang == "es" else "Document", "Versión" if lang == "es" else "Version", "Audiencia" if lang == "es" else "Audience"],
        [["ES" if lang == "es" else "EN", date.today().isoformat(), "Propietario o persona de soporte sin experiencia avanzada" if lang == "es" else "Owner or support person without advanced experience"]])
    note(doc, "Objetivo" if lang == "es" else "Purpose", "Dejar el sitio visible en Internet, mantenerlo actualizado y proteger la información personal." if lang == "es" else "Put the site online, keep it current and protect personal information.")
    doc.add_page_break()


def build(lang, filename):
    es = lang == "es"
    d = Document(); setup(d, lang); cover(d, lang)
    d.add_heading("1. Qué se entrega" if es else "1. What is included", 1)
    d.add_paragraph(("El portafolio es una aplicación web bilingüe, adaptable a móvil y escritorio, con modo claro y oscuro. Presenta el perfil profesional con lenguaje comprensible y una segunda capa técnica para reclutadores especializados.") if es else ("The portfolio is a bilingual web application for mobile and desktop with light and dark themes. It explains the professional profile in plain language and provides a deeper technical layer for specialist recruiters."))
    bullets(d, (["Web en español e inglés.", "Perfiles PDF en ambos idiomas dentro de una sección secundaria de documentos.", "Portada sin círculos ni coordenadas decorativas, enfocada en propuesta profesional y evidencia útil.", "Señal animada de cinco disciplinas con alternativa estática para usuarios que reducen movimiento.", "Emblemas angulares independientes para la línea digital y la línea automotriz, adaptables a ambos temas.", "Áreas independientes: ciberseguridad, infraestructura, software, asistencia técnica virtual y mecatrónica automotriz.", "Pruebas unitarias para preferencias, navegación y documentos, además de pruebas de renderizado.", "Logo FMV InfraSec adaptable al tema.", "Metadatos SEO, navegación accesible y preferencias persistentes."] if es else ["Spanish and English website.", "Profile PDFs in both languages inside a secondary documents section.", "Hero without decorative rings or coordinates, focused on professional value and useful proof.", "Animated five-discipline signal with a static alternative for reduced-motion users.", "Independent angular emblems for the digital and automotive tracks, adaptable to both themes.", "Independent areas: cybersecurity, infrastructure, software, technical virtual assistance and automotive mechatronics.", "Unit tests for preferences, navigation, and documents, plus rendered-output tests.", "Theme-aware FMV InfraSec logo.", "SEO metadata, accessible navigation and saved preferences."]))

    d.add_heading("2. Datos pendientes antes de publicar" if es else "2. Information to complete before launch", 1)
    comparison(d, ["Campo" if es else "Field", "Acción" if es else "Action", "Estado" if es else "Status"], [
        ["Teléfono corporativo" if es else "Corporate phone", "+57 302 528 6029 · WhatsApp", "Configurado" if es else "Configured"],
        ["Correo" if es else "Email", "ingenierossoftware31@gmail.com", "Configurado" if es else "Configured"],
        ["GitHub", "github.com/ingesoft04", "Configurado" if es else "Configured"],
        ["LinkedIn", "linkedin.com/in/anmartinez94", "Configurado" if es else "Configured"],
        ["Upwork", "freelancers/~0199b81b4f763382ef", "Configurado" if es else "Configured"],
        ["Workana", "Agregar el perfil cuando esté listo" if es else "Add profile when ready", "Opcional" if es else "Optional"],
        ["Dominio", "Elegir y comprar; ejemplo: fmv-infrasec.com" if es else "Choose and purchase; example: fmv-infrasec.com", "Opcional" if es else "Optional"],
    ])
    note(d, "Privacidad" if es else "Privacy", "No publique direcciones, documentos de identidad, credenciales, IP privadas/públicas, nombres de clientes ni detalles de infraestructura protegida." if es else "Do not publish addresses, identity documents, credentials, IP addresses, client names or protected infrastructure details.")

    d.add_heading("3. Requisitos del equipo" if es else "3. Computer requirements", 1)
    bullets(d, (["Windows 10/11, macOS o Linux.", "Node.js 22.13 o superior.", "pnpm instalado.", "Git instalado.", "Una cuenta de GitHub y, según la ruta elegida, una cuenta de Vercel o acceso a Codex Sites."] if es else ["Windows 10/11, macOS or Linux.", "Node.js 22.13 or newer.", "pnpm installed.", "Git installed.", "A GitHub account and either a Vercel account or Codex Sites access."]))
    code(d, "node --version\npnpm --version\ngit --version")
    d.add_paragraph(("Si alguno no responde con una versión, instálelo desde su sitio oficial y reinicie la terminal.") if es else ("If any command does not return a version, install that tool from its official website and restart the terminal."))

    d.add_heading("4. Estructura del proyecto" if es else "4. Project structure", 1)
    comparison(d, ["Ruta", "Contenido" if es else "Contents"], [
        ["app/page.tsx", "Texto, secciones, botones e interacción principal" if es else "Main content, sections, buttons and interactions"],
        ["app/components/", "Navegación, selector de oportunidades, proyectos y contacto" if es else "Navigation, opportunity selector, projects, and contact"],
        ["public/og.png", "Imagen social para LinkedIn, WhatsApp y otras plataformas" if es else "Social image for LinkedIn, WhatsApp, and other platforms"],
        ["app/hooks/", "Estado reutilizable para idioma y apariencia" if es else "Reusable state for language and appearance"],
        ["app/globals.css / recruiter.css", "Colores, componentes, adaptación móvil y temas" if es else "Colors, components, mobile layout and themes"],
        ["app/layout.tsx", "Título, descripción, palabras clave e icono" if es else "Title, description, keywords and icon"],
        ["public/", "Logo y archivos PDF descargables" if es else "Logo and downloadable PDF files"],
        ["build/", "Scripts para crear y validar documentos" if es else "Document generation and validation scripts"],
        [".openai/hosting.json", "Configuración del alojamiento con Sites" if es else "Sites hosting configuration"],
    ])
    note(d, "Arquitectura SOLID y responsive" if es else "SOLID and responsive architecture", "La página principal coordina; cada componente cumple una responsabilidad concreta y el hook administra únicamente preferencias. En escritorio la navegación es horizontal; a 900 píxeles o menos se convierte en un panel lateral accesible con botón, fondo de cierre y tecla Escape. La presentación también responde a 720, 580 y 420 píxeles." if es else "The main page coordinates; each component has one responsibility and the hook manages preferences only. Navigation is horizontal on desktop and becomes an accessible side drawer at 900 pixels or less, with a menu button, dismissible backdrop, and Escape-key support. The layout also responds at 720, 580, and 420 pixels.")

    d.add_heading("5. Ejecutar el portafolio localmente" if es else "5. Run the portfolio locally", 1)
    bullets(d, (["Abra PowerShell o Terminal.", "Entre en la carpeta del proyecto.", "Instale dependencias la primera vez.", "Inicie el servidor.", "Abra la dirección mostrada, normalmente http://localhost:3000/."] if es else ["Open PowerShell or Terminal.", "Change to the project folder.", "Install dependencies the first time.", "Start the development server.", "Open the displayed address, usually http://localhost:3000/."]), True)
    code(d, 'cd "C:\\Personal\\Portafolio"\npnpm install\npnpm run dev')
    note(d, "Detener" if es else "Stop", "En la terminal, presione Ctrl+C. Esto solo detiene la vista local; no borra archivos." if es else "Press Ctrl+C in the terminal. This only stops the local preview; it does not delete files.")

    d.add_heading("6. Personalizar contenido" if es else "6. Customize the content", 1)
    d.add_paragraph(("Edite app/page.tsx con Visual Studio Code. Busque el texto existente y cambie únicamente el valor entre comillas. Mantenga la estructura, comas y llaves. Para agregar el teléfono, use formato internacional (+57...) y un enlace https://wa.me/57NUMERO sin espacios.") if es else ("Edit app/page.tsx in Visual Studio Code. Search for existing text and change only the quoted value. Keep the surrounding structure, commas and braces. For the phone, use international format (+57...) and a https://wa.me/57NUMBER link without spaces."))
    bullets(d, (["Confirme cada cambio en español y en inglés.", "No invente métricas, clientes, certificaciones terminadas ni niveles de idioma.", "En proyectos confidenciales describa el problema, su responsabilidad, las tecnologías y el resultado general sin identificar al cliente.", "Pruebe todos los enlaces después de editarlos."] if es else ["Apply each change in both Spanish and English.", "Do not invent metrics, clients, completed certifications or language levels.", "For confidential projects, describe the problem, responsibility, technologies and general outcome without identifying the client.", "Test every link after editing."]))

    d.add_heading("7. Logo, PDFs y archivos públicos" if es else "7. Logo, PDFs and public files", 1)
    d.add_paragraph(("Conserve nombres simples, sin tildes ni espacios. Sustituya un archivo manteniendo exactamente el mismo nombre para evitar editar el código. Antes de reemplazar el logo, use PNG transparente y compruebe los modos claro y oscuro.") if es else ("Use simple names without accents or spaces. Replace a file using exactly the same name to avoid code changes. Before replacing the logo, use a transparent PNG and verify both light and dark modes."))
    code(d, "public/fmv-logo-white-transparent.png\npublic/fmv-logo-graphite-transparent.png\npublic/fmv-mark-color-transparent.png\npublic/og.png\npublic/perfil-andres-felipe-martinez-es.pdf\npublic/andres-felipe-martinez-profile-en.pdf")

    d.add_heading("8. Validación antes de publicar" if es else "8. Pre-deployment validation", 1)
    code(d, "pnpm run build")
    bullets(d, (["La compilación termina sin errores.", "Las pruebas unitarias y de renderizado terminan sin fallos.", "El selector ES/EN cambia todo el contenido.", "Los modos claro y oscuro mantienen contraste y muestran bien el logo.", "La señal de capacidades se adapta a móvil y respeta la preferencia de movimiento reducido.", "La portada no muestra círculos, coordenadas ni etiquetas decorativas.", "Los dos PDF están disponibles en la sección secundaria de documentos.", "La web se entiende sin conocimientos de software.", "Móvil: el menú lateral abre, cierra al elegir una sección y responde a la tecla Escape.", "Móvil: no hay texto cortado ni desplazamiento horizontal.", "Correo, GitHub, LinkedIn y WhatsApp abren el destino correcto."] if es else ["The build completes without errors.", "Unit and rendered-output tests complete without failures.", "The ES/EN selector changes all content.", "Light and dark themes have good contrast and display the logo correctly.", "The capability signal adapts to mobile and respects reduced-motion preferences.", "The hero shows no decorative rings, coordinates, or labels.", "Both PDFs are available from the secondary documents section.", "The site makes sense to a nontechnical reader.", "Mobile: the side menu opens, closes after choosing a section, and responds to Escape.", "Mobile: no clipped text or horizontal scrolling.", "Email, GitHub, LinkedIn and WhatsApp open the correct destination."]))

    d.add_heading("9. Elegir cómo publicarlo" if es else "9. Choose a publishing route", 1)
    comparison(d, ["Ruta" if es else "Route", "Ideal para" if es else "Best for", "Ventaja" if es else "Advantage", "Consideración" if es else "Consideration"], [
        ["Codex Sites", "Publicación directa desde este proyecto" if es else "Direct publishing from this project", "Flujo integrado" if es else "Integrated workflow", "Puede requerir autorización para hacerlo público" if es else "Public access may require authorization"],
        ["GitHub + Vercel", "Control propio y despliegues automáticos" if es else "Ownership and automatic deployments", "Cada cambio en GitHub se publica" if es else "Every GitHub update deploys", "Debe conectar dos cuentas" if es else "Two accounts must be connected"],
        ["Hosting propio", "Administración avanzada" if es else "Advanced administration", "Control total" if es else "Full control", "Más mantenimiento y seguridad" if es else "More maintenance and security work"],
    ])
    note(d, "Recomendación" if es else "Recommendation", "Use GitHub + Vercel para un portafolio profesional fácil de mantener; use Codex Sites cuando prefiera publicar y administrar desde Codex." if es else "Use GitHub + Vercel for an easy-to-maintain professional portfolio; use Codex Sites when you prefer to publish and manage it from Codex.")

    d.add_heading("10. Publicar con Codex Sites" if es else "10. Publish with Codex Sites", 1)
    bullets(d, ["Abra el proyecto en Codex.", "Solicite: “Valida y publica este portafolio con Sites”.", "Codex compila, empaqueta y crea una versión de alojamiento.", "Revise la URL de vista previa.", "Autorice explícitamente la publicación abierta si desea que cualquier reclutador pueda verla.", "Abra la URL final y complete la lista de validación."] if es else ["Open the project in Codex.", "Ask: “Validate and publish this portfolio with Sites.”", "Codex builds, packages and creates a hosting version.", "Review the preview URL.", "Explicitly approve open-world publishing if every recruiter should be able to access it.", "Open the final URL and complete the validation checklist."], True)
    d.add_paragraph(("El archivo .openai/hosting.json identifica el proyecto. No lo copie entre sitios distintos sin revisar el identificador.") if es else ("The .openai/hosting.json file identifies the project. Do not copy it between unrelated sites without checking the project identifier."))

    d.add_heading("11. Publicar con GitHub y Vercel" if es else "11. Publish with GitHub and Vercel", 1)
    bullets(d, ["Cree un repositorio privado o público vacío en GitHub; no agregue README desde GitHub.", "Desde la carpeta local, inicialice y suba el código.", "En Vercel, seleccione Add New Project, importe el repositorio y acepte la configuración detectada.", "Configure el comando de compilación como pnpm run build si no se detecta.", "Pulse Deploy y pruebe la URL asignada."] if es else ["Create an empty private or public GitHub repository; do not add a README on GitHub.", "Initialize and upload the code from the local folder.", "In Vercel, select Add New Project, import the repository and accept the detected configuration.", "Set the build command to pnpm run build if it is not detected.", "Select Deploy and test the assigned URL."], True)
    code(d, "git init\ngit add .\ngit commit -m \"Publicar portafolio inicial\"\ngit branch -M main\ngit remote add origin https://github.com/USUARIO/REPOSITORIO.git\ngit push -u origin main")
    note(d, "Repositorio privado" if es else "Private repository", "Vercel puede desplegarlo si autoriza el acceso. El código permanece privado, pero la web publicada puede ser pública." if es else "Vercel can deploy it when access is granted. The source remains private while the deployed website can be public.")

    d.add_heading("12. Conectar un dominio" if es else "12. Connect a custom domain", 1)
    bullets(d, ["Compre el dominio en un registrador reconocido.", "En Sites o Vercel, abra Domains y agregue el dominio.", "Copie exactamente los registros DNS solicitados (A, AAAA, CNAME o TXT).", "Espere la propagación: puede tardar desde minutos hasta 48 horas.", "Confirme HTTPS, el dominio con y sin www, y defina una redirección principal."] if es else ["Purchase the domain from a reputable registrar.", "In Sites or Vercel, open Domains and add the domain.", "Copy the requested DNS records exactly (A, AAAA, CNAME or TXT).", "Wait for DNS propagation, which may take minutes to 48 hours.", "Confirm HTTPS, both www and root variants, and choose one canonical redirect."], True)
    note(d, "Correo" if es else "Email", "No elimine registros MX o TXT existentes del correo corporativo. Cambie solo los registros indicados para la web." if es else "Do not delete existing MX or TXT records used by corporate email. Change only the records required for the website.")

    d.add_heading("13. Publicar actualizaciones" if es else "13. Publish updates", 1)
    code(d, "pnpm run build\ngit add .\ngit commit -m \"Actualizar experiencia y enlaces\"\ngit push")
    d.add_paragraph(("Con Vercel, el push activa una nueva publicación. Con Sites, solicite una nueva versión después de validar. Para revertir, seleccione una publicación anterior en el panel del proveedor; no borre el historial de Git.") if es else ("With Vercel, a push triggers a new deployment. With Sites, request a new version after validation. To roll back, select an earlier deployment in the provider dashboard; do not delete Git history."))

    d.add_heading("14. Seguridad y privacidad" if es else "14. Security and privacy", 1)
    bullets(d, (["Active autenticación multifactor en GitHub, Vercel, registrador y correo.", "No guarde contraseñas, tokens, llaves privadas ni archivos .env en Git.", "Use permisos mínimos y revoque accesos antiguos.", "No publique diagramas o datos que revelen defensas, IP, puertos o clientes.", "Mantenga dependencias actualizadas y revise alertas de seguridad.", "Haga copia del repositorio y conserve una versión local verificable.", "Revise trimestralmente enlaces, teléfono, tarifas, certificaciones y disponibilidad."] if es else ["Enable multi-factor authentication on GitHub, Vercel, the registrar and email.", "Never commit passwords, tokens, private keys or .env files.", "Use least-privilege access and revoke old integrations.", "Do not publish diagrams or data exposing defenses, IPs, ports or clients.", "Keep dependencies current and review security alerts.", "Back up the repository and keep a verifiable local copy.", "Review links, phone, rates, certifications and availability every quarter."]))

    d.add_heading("15. SEO, accesibilidad y selección internacional" if es else "15. SEO, accessibility and international recruiting", 1)
    bullets(d, (["Mantenga un título profesional directo y palabras clave reales.", "Use inglés claro, sin traducciones literales confusas.", "Incluya resultados verificables cuando estén disponibles.", "Preserve contraste, navegación por teclado y textos alternativos.", "Compruebe que el primer bloque explique qué hace, para quién y cómo contactarlo.", "Actualice zona horaria, disponibilidad y modalidad remota."] if es else ["Keep a direct professional title and truthful keywords.", "Use clear English without awkward literal translations.", "Add verifiable outcomes when available.", "Preserve contrast, keyboard navigation and alternative text.", "Ensure the first section explains what you do, who you help and how to contact you.", "Keep timezone, availability and remote-work preferences current."]))
    note(d, "URL pública y vista social" if es else "Public URL and social preview", "En el proveedor configure NEXT_PUBLIC_SITE_URL con la URL final, por ejemplo https://portafolio.example.com. Esto mantiene correcta la URL canónica y la tarjeta og.png que aparece al compartir en LinkedIn o WhatsApp." if es else "Set NEXT_PUBLIC_SITE_URL in the hosting provider to the final URL, for example https://portfolio.example.com. This keeps the canonical URL and the og.png sharing card correct on LinkedIn and WhatsApp.")

    d.add_heading("16. Solución de problemas" if es else "16. Troubleshooting", 1)
    comparison(d, ["Síntoma" if es else "Symptom", "Causa probable" if es else "Likely cause", "Solución" if es else "Fix"], [
        ["pnpm no se reconoce" if es else "pnpm is not recognized", "No está instalado o la terminal no se reinició" if es else "Not installed or terminal not restarted", "Instale pnpm y abra otra terminal" if es else "Install pnpm and open a new terminal"],
        ["Puerto 3000 ocupado" if es else "Port 3000 in use", "Otro servidor está activo" if es else "Another server is running", "Deténgalo con Ctrl+C o use la URL alternativa mostrada" if es else "Stop it with Ctrl+C or use the displayed alternative URL"],
        ["Build falla" if es else "Build fails", "Error de sintaxis o dependencia" if es else "Syntax or dependency error", "Revise el primer error, deshaga el último cambio y ejecute pnpm install" if es else "Read the first error, undo the last edit and run pnpm install"],
        ["PDF no descarga" if es else "PDF will not download", "Nombre o ruta no coincide" if es else "File name or path mismatch", "Verifique public/ y mayúsculas/minúsculas" if es else "Check public/ and letter case"],
        ["Dominio no abre" if es else "Domain does not open", "DNS pendiente o registro incorrecto" if es else "DNS pending or incorrect record", "Compare registros y espere propagación" if es else "Compare records and wait for propagation"],
        ["Logo se ve mal" if es else "Logo looks wrong", "Imagen sin transparencia o contraste insuficiente" if es else "No transparency or poor contrast", "Use PNG transparente y pruebe ambos temas" if es else "Use a transparent PNG and test both themes"],
    ])

    d.add_heading("17. Lista final de lanzamiento" if es else "17. Final launch checklist", 1)
    bullets(d, (["[ ] Teléfono y enlaces reales agregados", "[ ] Contenido ES/EN revisado", "[ ] Certificaciones expresadas como cursando cuando corresponda", "[ ] Build exitoso", "[ ] Revisión móvil y escritorio", "[ ] Modo claro y oscuro", "[ ] PDFs correctos", "[ ] HTTPS activo", "[ ] Dominio y redirección probados", "[ ] MFA y recuperación de cuentas configurados", "[ ] Sin secretos ni información confidencial", "[ ] Una persona no técnica entiende la propuesta en menos de un minuto"] if es else ["[ ] Real phone number and links added", "[ ] ES/EN content reviewed", "[ ] In-progress certifications labelled accurately", "[ ] Successful build", "[ ] Mobile and desktop review", "[ ] Light and dark mode checked", "[ ] Correct PDFs", "[ ] HTTPS active", "[ ] Domain and redirect tested", "[ ] MFA and account recovery configured", "[ ] No secrets or confidential information", "[ ] A nontechnical reader understands the offer in under one minute"]))

    d.add_heading("18. Glosario breve" if es else "18. Short glossary", 1)
    comparison(d, ["Término" if es else "Term", "Significado" if es else "Meaning"], [
        ["Build", "Versión optimizada del sitio lista para publicar" if es else "Optimized version of the site ready to publish"],
        ["Deploy", "Publicar una versión en un servidor" if es else "Publish a version to a server"],
        ["DNS", "Registros que conectan el dominio con el alojamiento" if es else "Records connecting the domain to hosting"],
        ["HTTPS", "Conexión cifrada indicada por el candado del navegador" if es else "Encrypted browser connection shown by a padlock"],
        ["Repositorio", "Historial organizado del código y sus cambios" if es else "Organized history of source code and changes"],
        ["Rollback", "Volver a una publicación anterior estable" if es else "Return to a previous stable deployment"],
    ])
    d.add_heading("19. Despliegue en Windows, Linux y Docker" if es else "19. Deployment on Windows, Linux, and Docker", 1)
    d.add_heading("Windows" if es else "Windows", 2)
    d.add_paragraph(("Use PowerShell con Node.js 22 y pnpm. Esta ruta sirve para desarrollo, pruebas o un servidor Windows administrado.") if es else ("Use PowerShell with Node.js 22 and pnpm. This route works for development, testing, or a managed Windows server."))
    code(d, 'cd "C:\\Personal\\Portafolio"\npnpm install\npnpm run build\npnpm start')
    d.add_heading("Linux", 2)
    d.add_paragraph(("Instale Node.js 22 y pnpm, compile el proyecto y mantenga el proceso con systemd o un administrador equivalente. Coloque Nginx, Caddy o Traefik delante del servicio para HTTPS.") if es else ("Install Node.js 22 and pnpm, build the project, and manage the process with systemd or an equivalent service manager. Put Nginx, Caddy, or Traefik in front of it for HTTPS."))
    code(d, "pnpm install\npnpm run build\nPORT=3000 pnpm start")
    d.add_heading("Docker Compose · Windows o Linux" if es else "Docker Compose · Windows or Linux", 2)
    d.add_paragraph(("Es la ruta más reproducible. Utiliza una imagen Node.js 22, comprobación de salud, sistema de archivos de solo lectura y capacidades Linux reducidas.") if es else ("This is the most reproducible route. It uses a Node.js 22 image, a health check, a read-only filesystem, and reduced Linux capabilities."))
    code(d, "docker compose up -d --build\ndocker compose ps\ndocker compose logs -f portfolio\ndocker compose down")
    note(d, "Puerto alternativo" if es else "Alternative port", "$env:PORTFOLIO_PORT=8080 en PowerShell; PORTFOLIO_PORT=8080 en Linux. Después abra http://localhost:8080/." if es else "Use $env:PORTFOLIO_PORT=8080 in PowerShell or PORTFOLIO_PORT=8080 on Linux. Then open http://localhost:8080/.")

    d.add_heading("20. Regenerar documentos después de cada cambio" if es else "20. Regenerate documents after every change", 1)
    bullets(d, (["Guarde los cambios de contenido o diseño.", "Regenerar los perfiles PDF en español e inglés.", "Regenerar ambos manuales DOCX cuando cambien instrucciones, contacto o despliegue.", "Revisar visualmente todas las páginas de los documentos.", "Ejecutar la compilación final.", "Publicar únicamente después de superar todas las validaciones."] if es else ["Save content or design changes.", "Regenerate the Spanish and English profile PDFs.", "Regenerate both DOCX manuals whenever instructions, contact details, or deployment change.", "Visually inspect every document page.", "Run the final production build.", "Publish only after all checks pass."]), True)
    code(d, "pnpm run docs:profiles\npnpm run build")
    note(d, "Regla del proyecto" if es else "Project rule", "Cada entrega futura debe incluir documentos renovados y una compilación exitosa, incluso cuando el cambio parezca pequeño." if es else "Every future delivery must include refreshed documents and a successful build, even when the change appears minor.")
    d.add_paragraph(("Documento preparado para FMV InfraSec. Actualícelo cuando cambien el proveedor de alojamiento, el dominio, los datos de contacto o la estructura del proyecto.") if es else ("Prepared for FMV InfraSec. Update this document whenever hosting, domain, contact details or project structure changes."))
    path = OUT / filename
    OUT.mkdir(parents=True, exist_ok=True)
    d.save(path)
    return path


if __name__ == "__main__":
    print(build("es", "Manual_Implementacion_Portafolio_FMV_ES_Actualizado.docx"))
    print(build("en", "FMV_Portfolio_Implementation_Manual_EN_Updated.docx"))
