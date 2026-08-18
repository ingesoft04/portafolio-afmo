from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "carta"
OUT.mkdir(parents=True, exist_ok=True)

NAVY = RGBColor(10, 42, 78)
CYAN = RGBColor(0, 156, 184)
INK = RGBColor(28, 33, 39)
MUTED = RGBColor(91, 99, 108)

DATA = {
    "es": {
        "file": "Carta_Presentacion_Andres_Felipe_Martinez_ES.docx",
        "label": "CARTA DE PRESENTACIÓN",
        "role": "Ingeniero de Software · Seguridad de la Información · Infraestructura",
        "date": "Bogotá, Colombia · [FECHA]",
        "recipient": "[EMPRESA]\nProceso: [CARGO U OPORTUNIDAD]",
        "greeting": "Estimado equipo de Selección:",
        "paras": [
            "Me interesa participar en el proceso para [CARGO U OPORTUNIDAD]. Soy Ingeniero de Software y especialista en formación en Seguridad de la Información, con experiencia en administración de servidores externos, soporte a clientes, implementación técnica, bases de datos SQL y documentación de incidentes.",
            "Puedo aportar una combinación práctica de software, seguridad e infraestructura. En mi trabajo actual acompaño la operación de servidores y la atención técnica a clientes; además, mantengo un laboratorio propio con Linux, máquinas virtuales, OpenMediaVault, Docker y Windows Server 2025 para validar configuraciones fuera de producción.",
            "Mi perfil también incluye formación y experiencia en mecatrónica automotriz. Esta trayectoria me ha dado una forma de trabajo orientada al diagnóstico: entender el problema, identificar riesgos, documentar hallazgos y proponer acciones claras para personas técnicas y no técnicas.",
            "Actualmente curso formación de CertiProf como Líder Implementador, Auditor Interno y Auditor Externo ISO/IEC 27001. Busco oportunidades remotas de tiempo completo, medio tiempo, por horas o por objetivos, tanto en equipos internacionales como en proyectos freelance.",
            "Agradezco su tiempo y la oportunidad de conversar sobre cómo puedo contribuir a [EMPRESA]. Mi portafolio y perfiles profesionales complementan esta carta con evidencia técnica, experiencia y datos de contacto.",
        ],
        "closing": "Atentamente,",
        "note": "Antes de enviar: reemplace los campos entre corchetes y adapte el primer y último párrafo a la vacante.",
    },
    "en": {
        "file": "Andres_Felipe_Martinez_Cover_Letter_EN.docx",
        "label": "COVER LETTER",
        "role": "Software Engineer · Information Security · Infrastructure",
        "date": "Bogotá, Colombia · [DATE]",
        "recipient": "[COMPANY]\nApplication: [ROLE OR OPPORTUNITY]",
        "greeting": "Dear Hiring Team,",
        "paras": [
            "I am interested in the [ROLE OR OPPORTUNITY] position. I am a Software Engineer pursuing a specialization in Information Security, with experience in external server administration, client support, technical implementation, SQL databases, and incident documentation.",
            "I bring a practical combination of software, security, and infrastructure skills. In my current role, I support server operations and client-facing technical work. I also maintain a home lab with Linux, virtual machines, OpenMediaVault, Docker, and Windows Server 2025 to validate configurations outside production.",
            "My background also includes training and hands-on experience in automotive mechatronics. This path shaped a diagnostic approach to work: understand the problem, identify risks, document findings, and communicate clear actions to technical and nontechnical stakeholders.",
            "I am currently completing CertiProf training as an ISO/IEC 27001 Lead Implementer, Internal Auditor, and External Auditor. I am open to remote full-time, part-time, hourly, or outcome-based work with international teams and freelance clients.",
            "Thank you for your time and consideration. I would welcome a conversation about how I can contribute to [COMPANY]. My portfolio and professional profiles provide additional evidence of my technical work, experience, and contact information.",
        ],
        "closing": "Sincerely,",
        "note": "Before sending: replace every bracketed field and tailor the first and final paragraphs to the role.",
    },
}

def font(run, size=11, bold=False, color=INK, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size); run.bold = bold; run.italic = italic; run.font.color.rgb = color

def shade_paragraph(paragraph, fill):
    ppr = paragraph._p.get_or_add_pPr(); shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), fill); ppr.append(shd)

def add_link(paragraph, text, url):
    part = paragraph.part; rid = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    link = OxmlElement("w:hyperlink"); link.set(qn("r:id"), rid); run = OxmlElement("w:r"); rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color"); color.set(qn("w:val"), "006E86"); rpr.append(color); underline = OxmlElement("w:u"); underline.set(qn("w:val"), "single"); rpr.append(underline)
    text_el = OxmlElement("w:t"); text_el.text = text; run.append(rpr); run.append(text_el); link.append(run); paragraph._p.append(link)

def create(lang):
    d = DATA[lang]; doc = Document(); section = doc.sections[0]
    section.page_width = Inches(8.5); section.page_height = Inches(11)
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(.492)
    normal = doc.styles["Normal"]; normal.font.name = "Calibri"; normal.font.size = Pt(11); normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0); normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.10

    header = section.header.paragraphs[0]; header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font(header.add_run("FMV InfraSec  |  SOFTWARE · SECURITY · INFRASTRUCTURE"), 8.5, True, MUTED)
    footer = section.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(footer.add_run("Andrés Felipe Martínez Obando · ingenierossoftware31@gmail.com · +57 302 528 6029"), 8.5, False, MUTED)

    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    font(p.add_run(d["label"]), 9, True, CYAN)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(3)
    font(p.add_run("Andrés Felipe Martínez Obando"), 23, True, NAVY)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(10)
    font(p.add_run(d["role"]), 12.5, False, MUTED)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(18)
    font(p.add_run(d["date"]), 10, False, MUTED)

    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(14)
    recipient_lines = d["recipient"].split("\n")
    for i, line in enumerate(recipient_lines):
        font(p.add_run(line + ("\n" if i < len(recipient_lines) - 1 else "")), 10.5, i == 0, INK)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(10); font(p.add_run(d["greeting"]), 11, True, INK)
    for text in d["paras"]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.space_after = Pt(8); p.paragraph_format.line_spacing = 1.10; font(p.add_run(text), 10.5)
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(4); font(p.add_run(d["closing"]), 10.5)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2); font(p.add_run("Andrés Felipe Martínez Obando"), 11, True, NAVY)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2); font(p.add_run("Bogotá, Colombia · UTC-5"), 9.5, False, MUTED)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(8)
    add_link(p, "LinkedIn", "https://www.linkedin.com/in/anmartinez94"); font(p.add_run("  ·  "), 9.5, False, MUTED); add_link(p, "GitHub", "https://github.com/ingesoft04"); font(p.add_run("  ·  "), 9.5, False, MUTED); add_link(p, "Upwork", "https://www.upwork.com/freelancers/~0199b81b4f763382ef")
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(0); shade_paragraph(p, "EAF7F9"); font(p.add_run("  " + d["note"] + "  "), 8.5, False, MUTED, True)
    path = OUT / d["file"]; doc.save(path); print(path)

for language in DATA:
    create(language)
